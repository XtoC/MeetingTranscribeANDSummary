from docx import Document

def transcribe(filename, model):
    segments, _ = model.transcribe(filename, beam_size=5)

    return segments

def save_transcript(segments, filename):
    doc = Document()
    doc.add_heading("Meeting Transcript", level=1)
    text = ""
    for seg in segments:
        # print(seg.text)
        doc.add_paragraph(seg.text)
        text += seg.text
    doc.save(filename)
    print(f"Transcript saved to {filename}")

    return text