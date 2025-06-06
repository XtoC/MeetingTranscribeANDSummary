from gpt4all import GPT4All
from docx import Document
from transformers import pipeline
def chunk_text(text, max_chunk_size=350):
    """
    Split text into chunks roughly max_chunk_size tokens (approx words here).
    This naive split just cuts by words.
    """
    words = text.split()
    chunks = []
    for i in range(0, len(words), max_chunk_size):
        chunk = " ".join(words[i:i + max_chunk_size])
        chunks.append(chunk)
    return chunks


def summarize_chunk(model, chunk, max_tokens=1000):
    prompt = (
            "Please provide a concise summary of the following meeting transcript section:\n\n"
            + chunk
            + "\n\nSummary:"
    )
    summary = model.generate(prompt, max_tokens=max_tokens)
    return summary.strip()


def summarize_long_transcript(transcript, model_path, chunk_size=1000, max_tokens=1000):
    # Load model
    gptj = GPT4All(model_name="DeepSeek-R1-Distill-Qwen-7B~Q4_0.gguf", model_path=model_path, allow_download=False)

    chunks = chunk_text(transcript, max_chunk_size=chunk_size)
    chunk_summaries = []

    for i, chunk in enumerate(chunks):
        print(f"Summarizing chunk {i + 1}/{len(chunks)}...")
        summary = summarize_chunk(gptj, chunk, max_tokens=max_tokens)
        chunk_summaries.append(summary)

    # Optionally combine all chunk summaries into a final summary
    # final_prompt = (
    #         "The following are summaries of sections from a meeting transcript. "
    #         "Please provide an overall concise summary of the entire meeting:\n\n"
    #         + "\n\n".join(chunk_summaries)
    #         + "\n\nFinal Summary:"
    # )
    # final_summary = gptj.generate(final_prompt, max_tokens=max_tokens)

    return chunk_summaries

def summarize_by_model(input_text, model, chunk_max_token=512, max_length = 160, min_length=60):
    summarizer = pipeline("summarization", model=model)
    chunk_transcript = chunk_text(input_text, chunk_max_token)
    chunk_summaries = []
    for transcript in chunk_transcript:
        summary = summarizer(transcript, max_length=max_length, min_length=min_length, do_sample=False)
        chunk_summaries.append(summary[0]['summary_text'])

    return chunk_summaries

def save_summary(filename, chunk_summaries):
    doc = Document()
    doc.add_heading("Meeting Summary", level=1)
    for chunk in chunk_summaries:
        doc.add_paragraph(chunk)
    doc.save(filename)
    print(f"Summary saved to {filename}")